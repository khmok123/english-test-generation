from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from shutil import copyfile
from tense_exercise import tense_exercise_auto
from vocab_exercise import vocab_exercise_auto
from reading_exercise import reading_exercise_auto
from prepositions_exercise import prepositions_exercise_auto
from adj_exercise import adj_exercise_auto
from doc_func import add_test_title, add_one_line, add_section_title, add_allowed_time, add_regular_text, add_line_break, add_page_break

template_dir = r'.\template\template.docx'
output_dir = r'.\template\output.docx'

# -------------------------------------------------------------------------------------
alignments = {"center": WD_ALIGN_PARAGRAPH.CENTER, 
              "left": WD_ALIGN_PARAGRAPH.LEFT,
             "right": WD_ALIGN_PARAGRAPH.RIGHT}

test_title_style = {"fontname": "Arial", "fontsize": Pt(16), 
                "alignment": alignments["center"]}
allowed_time_style = {"fontname": "Arial", "fontsize": Pt(12), 
               "alignment": alignments["left"], 
                "bold": False, "italic": False, "underline": False}
section_title_style = {"fontname": "Arial", "fontsize": Pt(12),
                      "alignment": alignments["left"],
                "bold": True, "italic": False, "underline": False}
regular_text_style = {"fontname": "Arial", "fontsize": Pt(12), 
               "alignment": alignments["left"], 
                "bold": False, "italic": False, "underline": False}
#-----------------------------------------------------------------------------------
def run_grammar_sections():
    sections = ['Tense', 'Vocabulary', 'Prepositions', 'Adjectives and adverbs']
    print("Grammar sections:")
    for idx, section in enumerate(sections):
        print(str(idx+1)+':', section)
    section_idx_str = input("Enter section indices (seperated with commas): ")
    section_idx_list = [int(idx.strip()) for idx in section_idx_str.split(',')]
    for idx in section_idx_list:
        if idx==1:
            tense_exercise_auto(document,grade)
        elif idx==2:
            vocab_exercise_auto(document)
        elif idx==3:
            prepositions_exercise_auto(document)
        elif idx==4:
            adj_exercise_auto(document)
#-------------------------------------------------------------------------------------

copyfile(template_dir, output_dir)
document = Document(output_dir)
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
# for section in document.sections:
#     first_page_footer = section.first_page_footer
#     footer = section.footer
#     first_page_footer.paragraphs[0].text  = "testing"
#     footer.paragraphs[0].text  = "testing"
    
print("Welcome to the English test generator!")
grade = int(input("Enter the grade of the test: "))
test_title = input("Enter the title of the test: ")
unit_title = input("Enter the unit title of the test: ")
add_test_title(document, test_title_style, test_title, unit_title)

add_allowed_time(document, "Time allowed: 45 minutes")
add_section_title(document, "A.  Grammar and Usage (60%)")

run_grammar_sections()

# tense_exercise_auto(document,grade)
# vocab_exercise_auto(document)
# prepositions_exercise_auto(document)
# adj_exercise_auto(document)

reading_exercise_auto(document)


document.save(output_dir)



    



