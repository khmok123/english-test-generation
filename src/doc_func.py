from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import re

alignments = {"center": WD_ALIGN_PARAGRAPH.CENTER, 
              "left": WD_ALIGN_PARAGRAPH.LEFT,
             "right": WD_ALIGN_PARAGRAPH.RIGHT}

test_title_style = {"fontname": "Arial", "fontsize": Pt(16), 
                "alignment": alignments["center"]}
allowed_time_style = {"fontname": "Arial", "fontsize": Pt(12), 
               "alignment": alignments["left"], 
                "bold": False, "italic": False, "underline": False}
section_title_style = {"fontname": "Arial", "fontsize": Pt(12),
                      "alignment": alignments["left"],
                "bold": True, "italic": False, "underline": False}
regular_text_style = {"fontname": "Arial", "fontsize": Pt(12), 
               "alignment": alignments["left"], 
                "bold": False, "italic": False, "underline": False}

def add_test_title(document, style, test_title, unit_name):
    title = document.paragraphs[0]
    title.alignment = style["alignment"]
    run1 = title.add_run(test_title)
    run1.font.name = style["fontname"]
    run1.font.size = style["fontsize"]
    run1.underline = True
    run1.bold = True
    run1.add_break()
    run2 = title.add_run(unit_name)
    run2.font.name = style["fontname"]
    run2.font.size = style["fontsize"]
    run2.bold = True
    
def add_one_line(document, style, text):
    p = document.add_paragraph()
    p.alignment = style["alignment"]
    run = p.add_run(text)
    run.font.name = style["fontname"]
    run.font.size = style["fontsize"]
    run.bold = style["bold"]
    run.italic = style["italic"]
    run.underline = style["underline"]
    
def add_to_cell(table, text):
    style = regular_text_style
    cell = table.cell(0, 0)
    paragraphs = cell.paragraphs
    paragraphs[0].alignment = alignments["center"]
    run = paragraphs[0].add_run(text)
    run.font.name = style["fontname"]
    run.font.size = style["fontsize"]
    run.bold = style["bold"]
    run.italic = style["italic"]
    run.underline = style["underline"]    

def add_section_title(document, text):
    add_one_line(document, section_title_style, text)
    
def add_allowed_time(document, text):
    add_one_line(document, allowed_time_style, text)
    
def add_regular_text(document, text):
    add_one_line(document, regular_text_style, text)
    
def add_line_break(document):
    p = document.add_paragraph()
    run = p.add_run('')
    run.add_break(WD_BREAK.LINE)
    
def add_page_break(document):
    p = document.add_paragraph()
    run = p.add_run('')
    run.add_break(WD_BREAK.PAGE)
    
    
def untokenize(words):
    """
    Untokenizing a text undoes the tokenizing operation, restoring
    punctuation and spaces to the places that people expect them to be.
    Ideally, `untokenize(tokenize(text))` should be identical to `text`,
    except for line breaks.
    """
    text = ' '.join(words)
    step1 = text.replace("`` ", '"').replace(" ''", '"').replace('. . .',  '...')
    step2 = step1.replace(" ( ", " (").replace(" ) ", ") ")
    step3 = re.sub(r' ([.,:;?!%]+)([ \'"`])', r"\1\2", step2)
    step4 = re.sub(r' ([.,:;?!%]+)$', r"\1", step3)
    step5 = step4.replace(" '", "'").replace(" n't", "n't").replace(
         "can not", "cannot")
    step6 = step5.replace(" ` ", " '")
    return step6.strip()

    

# def add_ordered_list(document, str_list):
#     style = regular_text_style
#     for element in str_list:
#         p = document.add_paragraph('')
# #         p.style = "List Number"
#         p.alignment = style["alignment"]
#         run = p.add_run(element)
#         run.font.name = style["fontname"]
#         run.font.size = style["fontsize"]
#         run.bold = style["bold"]
#         run.italic = style["italic"]
#         run.underline = style["underline"]
    